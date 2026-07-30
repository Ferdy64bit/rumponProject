from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\ThinkPad T14\Documents\aplikasi java\java3")
OUT = Path(r"C:\Users\ThinkPad T14\Music\BAB IV Lanjutan Fishing Point.docx")
IMG = ROOT / "app/src/main/java/com/example/java3/readme/screenshots"


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.57)
section.bottom_margin = Inches(1.18)
section.left_margin = Inches(1.57)
section.right_margin = Inches(1.18)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    if style_name != "Normal":
        style.font.bold = True

counter = {"gambar": 1, "tabel": 1}


def add_paragraph(text="", bold=False, align=None, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Inches(0.39) if indent else Inches(0)
    para.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    return para


def heading(text, level=2):
    para = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    para.paragraph_format.first_line_indent = Inches(0)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = True


def add_code(text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.left_indent = Inches(0.35)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10)


def add_table(caption, headers, rows):
    add_paragraph(f"Tabel 4.{counter['tabel']} {caption}", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    counter["tabel"] += 1
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for idx, head in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        cell.text = str(head)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[idx].paragraphs:
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    add_paragraph("", indent=False)


def add_image(filename, caption, width=4.2):
    path = IMG / filename
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Inches(0)
        para.add_run().add_picture(str(path), width=Inches(width))
    else:
        add_paragraph(f"[Gambar Screenshot {caption}]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(f"Gambar 4.{counter['gambar']} {caption}", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    counter["gambar"] += 1


def add_intro():
    add_paragraph("BAB IV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph("HASIL DAN PEMBAHASAN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph("")
    heading("4.3 Implementasi Sistem")
    add_paragraph("Tahap implementasi sistem merupakan proses penerjemahan rancangan aplikasi Fishing Point ke dalam bentuk perangkat lunak Android yang dapat digunakan oleh pengguna. Pada tahap ini, rancangan kebutuhan fungsional, rancangan basis data, rancangan antarmuka, dan rancangan Recommendation Engine diimplementasikan menggunakan Android Studio dengan bahasa pemrograman Java. Implementasi dilakukan dengan tetap mengikuti pola arsitektur MVVM dan Repository Pattern agar tanggung jawab antar komponen dapat dipisahkan secara jelas.")
    add_paragraph("Secara umum, struktur kode aplikasi dibagi ke dalam beberapa lapisan, yaitu core, data, domain, dan presentation. Lapisan core berisi utilitas umum seperti pengelolaan koneksi, session, konstanta, dan fungsi lokasi. Lapisan data berisi model response API serta repository yang menghubungkan aplikasi dengan Firebase, Cloud Firestore, Cloudinary, OpenWeather API, Open-Meteo Marine API, dan BMKG. Lapisan domain digunakan untuk menyimpan model bisnis serta Recommendation Engine, sedangkan lapisan presentation berisi Activity, Fragment, Adapter, dan ViewModel yang berhubungan langsung dengan tampilan aplikasi.")
    add_paragraph("Pemisahan tersebut membuat alur data lebih terarah. Pengguna berinteraksi melalui Activity atau Fragment, kemudian ViewModel meneruskan permintaan ke Repository. Repository mengambil data dari Firebase atau API eksternal, lalu hasilnya dikirim kembali ke ViewModel untuk ditampilkan pada antarmuka. Jika data memerlukan proses perhitungan, seperti jarak dan rekomendasi, maka data tersebut diproses terlebih dahulu pada lapisan domain sebelum ditampilkan kepada pengguna.")
    add_code("Activity / Fragment\n        ↓\nViewModel\n        ↓\nRepository\n        ↓\nFirebase / API Eksternal\n        ↓\nDomain Service (Haversine dan Recommendation Engine)\n        ↓\nViewModel\n        ↓\nTampilan Aplikasi")
    add_paragraph("Diagram alur tersebut menunjukkan bahwa aplikasi tidak mengambil dan memproses data secara langsung pada tampilan. Dengan demikian, implementasi menjadi lebih terstruktur dan lebih mudah diuji karena logika bisnis tidak bercampur dengan komponen antarmuka.")


def add_auth_dashboard_maps():
    heading("4.3.1 Authentication", 3)
    add_paragraph("Fitur authentication diimplementasikan untuk mengatur akses pengguna terhadap aplikasi. Pengguna harus melakukan login atau register sebelum dapat menggunakan fitur utama seperti membuat fishing point, menyimpan favorite, membuat postingan komunitas, dan mengelola profil. Implementasi authentication menggunakan Firebase Authentication karena layanan tersebut mendukung autentikasi berbasis email dan password serta menyediakan fitur reset password.")
    add_paragraph("Komponen utama yang digunakan pada fitur ini adalah SplashActivity, LoginActivity, RegisterActivity, AuthRepository, SessionManager, dan FirebaseManager. SplashActivity bertugas memeriksa status session ketika aplikasi pertama kali dijalankan. Apabila session pengguna masih tersedia, aplikasi akan langsung membuka MainActivity. Jika session tidak ditemukan, aplikasi menampilkan halaman login. LoginActivity dan RegisterActivity digunakan sebagai halaman input email dan password, sedangkan AuthRepository menjadi penghubung antara tampilan aplikasi dan Firebase Authentication.")
    add_paragraph("Alur login dimulai ketika pengguna memasukkan email dan password. Data tersebut diteruskan dari Activity menuju ViewModel, kemudian ViewModel memanggil AuthRepository. Repository melakukan proses autentikasi menggunakan Firebase Authentication. Jika proses berhasil, SessionManager menyimpan status login pengguna dan aplikasi membuka halaman utama. Apabila gagal, pesan kesalahan dikembalikan ke tampilan agar pengguna dapat memperbaiki data yang dimasukkan.")
    add_image("runtime-audit-current.png", "Tampilan aplikasi setelah proses authentication berhasil")
    add_paragraph("Gambar tersebut memperlihatkan bahwa pengguna yang telah berhasil melewati proses authentication dapat mengakses halaman utama aplikasi. Hal ini menunjukkan bahwa alur session berjalan dengan baik, karena aplikasi tidak kembali ke halaman login setelah pengguna dinyatakan aktif.")
    add_paragraph("Fitur reset password juga disediakan untuk membantu pengguna yang lupa kata sandi. Proses reset password memanfaatkan fungsi reset password dari Firebase Authentication, sehingga sistem akan mengirimkan tautan reset ke email pengguna yang terdaftar. Pada aplikasi, pengguna diberi informasi bahwa email reset dapat masuk ke folder spam agar proses pemulihan akun tidak dianggap gagal.")
    heading("4.3.2 Dashboard", 3)
    add_paragraph("Dashboard merupakan halaman utama yang menampilkan ringkasan kondisi memancing. Halaman ini dirancang agar pengguna dapat memperoleh informasi penting tanpa harus membuka banyak halaman. Informasi yang ditampilkan meliputi lokasi pengguna, kondisi cuaca live, kondisi perairan, grafik gelombang, safety score, fish activity, rekomendasi memancing, serta daftar spot terdekat.")
    add_paragraph("Implementasi dashboard dilakukan melalui HomeFragment dan HomeViewModel. HomeFragment bertanggung jawab menampilkan data ke antarmuka, sedangkan HomeViewModel bertugas mengambil dan mengolah data. Data cuaca diperoleh melalui WeatherRepository yang terhubung dengan OpenWeather API. Data perairan diperoleh melalui MarineHourlyRepository dan TideRepository yang mengakses Open-Meteo Marine API serta BMKG. Data fishing point diperoleh dari FishingRepository yang mengambil data dari Cloud Firestore.")
    add_paragraph("Setelah seluruh data diperoleh, HomeViewModel menghitung jarak antara pengguna dan setiap spot menggunakan LocationUtils. Spot terdekat kemudian dikirim ke RecommendationEngine untuk memperoleh skor rekomendasi. Hasil perhitungan tersebut ditampilkan pada dashboard dalam bentuk ringkasan yang terdiri atas label rekomendasi, persentase skor, indikator safety, indikator aktivitas ikan, dan informasi kondisi perairan.")
    add_image("fishingpoint_dashboard_real_device_after_fix.png", "Tampilan dashboard aplikasi Fishing Point pada perangkat Android nyata")
    add_paragraph("Gambar tersebut menunjukkan tampilan dashboard setelah dilakukan penyesuaian antarmuka. Informasi utama dibuat dalam bentuk card agar mudah dibaca. Bagian rekomendasi menampilkan hasil perhitungan Recommendation Engine, sedangkan bagian kondisi perairan menampilkan data gelombang dan informasi maritim yang diperoleh dari API.")
    add_image("dashboard-ui-final-live.png", "Tampilan dashboard dengan data cuaca, perairan, dan rekomendasi live")
    add_paragraph("Gambar tersebut menunjukkan bahwa dashboard telah mengintegrasikan beberapa sumber data. OpenWeather digunakan untuk membaca kondisi cuaca aktual, Open-Meteo Marine digunakan untuk membaca gelombang dan arus per jam, sedangkan BMKG digunakan sebagai informasi prakiraan dan peringatan resmi. Kombinasi data tersebut membuat dashboard dapat memberikan informasi yang lebih kontekstual kepada pengguna.")
    heading("4.3.3 Maps", 3)
    add_paragraph("Modul Maps diimplementasikan untuk menampilkan lokasi pengguna dan lokasi fishing point dalam bentuk peta digital. Implementasi fitur ini menggunakan Google Maps API, Fused Location Provider, MapFragment, MapViewModel, FishingRepository, dan FishingMarkerRenderer. Google Maps API digunakan sebagai media visualisasi, sedangkan Fused Location Provider digunakan untuk memperoleh koordinat pengguna secara real-time.")
    add_paragraph("Data fishing point yang tersimpan di Cloud Firestore diambil melalui FishingRepository. Setelah data diterima, MapViewModel mengirimkan daftar spot ke MapFragment untuk ditampilkan sebagai marker pada peta. Marker berisi informasi dasar seperti nama spot, jenis spot, jarak dari pengguna, dan status akses. Aplikasi juga menerapkan aturan ownership sehingga fitur edit dan hapus hanya tersedia bagi pengguna yang membuat spot tersebut.")
    add_paragraph("Pada implementasi terbaru, aplikasi mendukung dua jenis spot, yaitu spot publik dan spot pribadi. Spot publik dapat dilihat oleh pengguna lain, sedangkan spot pribadi hanya dapat dilihat oleh pemiliknya. Pengguna juga dapat mengubah status akses spot melalui fitur edit. Pendekatan ini diterapkan karena tidak semua pemancing ingin membagikan lokasi pancing pribadinya kepada publik.")
    add_image("fishingpoint_maps_validation.png", "Tampilan Google Maps dengan marker fishing point")
    add_paragraph("Gambar tersebut memperlihatkan marker fishing point yang ditampilkan pada peta. Marker menjadi representasi visual dari data lokasi yang tersimpan pada Firestore. Dengan tampilan ini, pengguna dapat mengetahui persebaran titik pancing di sekitar wilayahnya.")
    add_image("fishingpoint_maps_spot_card_validation.png", "Tampilan informasi spot pada halaman maps")
    add_paragraph("Selain menampilkan marker, aplikasi juga menyediakan card informasi spot ketika salah satu marker dipilih. Card tersebut membantu pengguna melihat informasi awal sebelum membuka halaman detail. Pada wilayah perairan tertentu, rute Google Maps tidak selalu tersedia. Oleh karena itu, aplikasi menambahkan polyline sebagai garis bantu dari lokasi pengguna menuju spot tujuan.")


def add_detail_haversine():
    heading("4.3.4 Detail Fishing Point", 3)
    add_paragraph("Detail Fishing Point diimplementasikan melalui DetailSpotActivity. Halaman ini digunakan untuk menampilkan informasi lengkap mengenai spot yang dipilih pengguna. Informasi yang ditampilkan mencakup nama spot, jenis spot, foto, koordinat, pemilik spot, status favorite, tombol navigasi, tombol share, serta data cuaca dan perairan berdasarkan koordinat spot.")
    add_paragraph("Perbedaan penting antara dashboard dan detail spot terletak pada sumber koordinat data lingkungan. Dashboard menggunakan koordinat pengguna saat ini, sedangkan detail spot menggunakan koordinat spot yang dipilih. Dengan demikian, data cuaca dan kondisi perairan yang tampil pada detail spot lebih relevan terhadap lokasi tujuan. Hal ini penting karena kondisi perairan pada wilayah pesisir dapat berbeda antara satu titik dengan titik lainnya.")
    add_paragraph("Fitur favorite pada detail spot memungkinkan pengguna menyimpan lokasi yang dianggap penting. Data favorite disimpan pada Firestore melalui FavoriteRepository. Fitur share memungkinkan pengguna membagikan informasi spot melalui aplikasi lain. Bagi pemilik spot, halaman detail juga menyediakan fitur edit, hapus, dan pengelolaan foto spot. Foto spot dapat diunggah melalui Cloudinary. Jika foto dihapus, aplikasi mengembalikannya ke gambar default berdasarkan jenis spot.")
    add_image("fishingpoint_maps_spot_card_validation_2.png", "Tampilan validasi informasi detail fishing point")
    heading("4.3.5 Implementasi Metode Haversine", 3)
    add_paragraph("Metode Haversine diimplementasikan pada kelas LocationUtils. Fungsi ini digunakan untuk menghitung jarak antara koordinat pengguna dan koordinat fishing point. Hasil perhitungan jarak digunakan pada dashboard, maps, detail spot, daftar spot, dan Recommendation Engine. Dengan demikian, Haversine menjadi salah satu bagian utama dalam penelitian ini karena berhubungan langsung dengan tujuan aplikasi untuk menentukan lokasi memancing berdasarkan jarak.")
    add_paragraph("Potongan implementasi metode Haversine pada aplikasi ditampilkan sebagai berikut.")
    add_code("public static double calculateDistance(double lat1, double lon1, double lat2, double lon2) {\n    final int R = 6371;\n    double latDistance = Math.toRadians(lat2 - lat1);\n    double lonDistance = Math.toRadians(lon2 - lon1);\n    double a = Math.sin(latDistance / 2) * Math.sin(latDistance / 2)\n        + Math.cos(Math.toRadians(lat1)) * Math.cos(Math.toRadians(lat2))\n        * Math.sin(lonDistance / 2) * Math.sin(lonDistance / 2);\n    double c = 2 * Math.atan2(Math.sqrt(a), Math.sqrt(1 - a));\n    return R * c;\n}")
    add_paragraph("Pada source code tersebut, nilai latitude dan longitude terlebih dahulu dikonversi ke radian. Selanjutnya, sistem menghitung nilai a sebagai komponen utama rumus Haversine. Nilai c menunjukkan sudut pusat antara dua titik pada permukaan bumi, sedangkan hasil akhir diperoleh dengan mengalikan c dengan jari-jari bumi. Nilai yang dikembalikan oleh fungsi tersebut menggunakan satuan kilometer.")
    add_table("Penilaian Distance Score", ["No", "Jarak Spot dari Pengguna", "Distance Score"], [[1, "≤ 2 km", 100], [2, "> 2 km sampai ≤ 5 km", 90], [3, "> 5 km sampai ≤ 10 km", 75], [4, "> 10 km sampai ≤ 20 km", 60], [5, "> 20 km sampai ≤ 50 km", 40], [6, "> 50 km", 20]])
    add_paragraph("Tabel tersebut menunjukkan konversi jarak menjadi Distance Score. Semakin dekat spot dari pengguna, semakin tinggi nilai yang diberikan. Nilai ini kemudian digabungkan dengan komponen lain pada Recommendation Engine. Namun, aplikasi juga menerapkan batas skor akhir berdasarkan jarak agar spot yang terlalu jauh tidak memperoleh rekomendasi terlalu tinggi hanya karena faktor cuaca dan perairan sedang baik.")


def add_recommendation():
    heading("4.3.6 Implementasi Recommendation Engine", 3)
    add_paragraph("Recommendation Engine merupakan komponen inti pada aplikasi Fishing Point. Engine ini bertugas menghasilkan rekomendasi lokasi memancing berdasarkan gabungan beberapa parameter. Implementasi dilakukan pada kelas RecommendationEngine yang berada pada lapisan domain. Penempatan ini bertujuan agar perhitungan rekomendasi tidak bercampur dengan kode tampilan dan dapat diuji melalui unit testing.")
    add_paragraph("Data yang digunakan oleh Recommendation Engine berasal dari beberapa sumber. Distance Score diperoleh dari hasil Haversine. Weather Score diperoleh dari OpenWeather API. Marine Score diperoleh dari Open-Meteo Marine API. Fish Activity Score diperoleh dari kombinasi solunar, cuaca, tekanan udara, dan pergerakan air. Spot Quality Score diperoleh dari rating spot, sedangkan User Preference Score digunakan untuk menyesuaikan rekomendasi dengan kecenderungan pengguna. Setelah skor dasar terbentuk, Safety Engine digunakan sebagai faktor koreksi agar rekomendasi tetap mempertimbangkan keselamatan.")
    add_code("GPS Pengguna + Data Spot\n        ↓\nDistance Engine\n        ↓\nWeather Engine + Marine Engine\n        ↓\nFish Activity Engine\n        ↓\nSpot Quality + User Preference\n        ↓\nBase Recommendation Score\n        ↓\nSafety Engine\n        ↓\nFinal Recommendation")
    add_code("Base Score = (0.20 × D) + (0.20 × W) + (0.25 × M) + (0.15 × A) + (0.10 × S) + (0.10 × U)\nFinal Score = Base Score × Safety Factor")
    add_paragraph("Formula tersebut menunjukkan bahwa kondisi perairan memperoleh bobot paling besar karena aplikasi berfokus pada aktivitas memancing di wilayah pesisir. Distance Score dan Weather Score juga memiliki pengaruh penting, sedangkan Fish Activity, Spot Quality, dan User Preference digunakan sebagai faktor pendukung agar rekomendasi lebih kontekstual.")
    for title, text in [
        ("4.3.6.1 Distance Engine", "Distance Engine berfungsi mengubah jarak hasil Haversine menjadi skor internal. Input dari engine ini adalah jarak dalam kilometer, sedangkan output yang dihasilkan adalah Distance Score pada rentang 0 sampai 100. Semakin dekat spot dari pengguna, semakin tinggi skor yang dihasilkan. Selain itu, sistem menerapkan distance score cap agar rekomendasi tidak terlalu tinggi untuk spot yang sangat jauh."),
        ("4.3.6.2 Weather Engine", "Weather Engine membaca data cuaca live dari OpenWeather API. Parameter yang digunakan meliputi kondisi langit, deskripsi cuaca, kecepatan angin, tekanan udara, dan jarak pandang. Pada proses perhitungan, kondisi cerah dan berawan memperoleh nilai lebih tinggi, sedangkan hujan lebat dan badai petir memperoleh nilai rendah karena berisiko terhadap kenyamanan dan keselamatan pengguna."),
        ("4.3.6.3 Marine Engine", "Marine Engine membaca kondisi perairan dari Open-Meteo Marine API. Data yang digunakan adalah tinggi gelombang saat ini, tinggi gelombang maksimum 24 jam ke depan, stabilitas gelombang, suhu permukaan laut, dan kecepatan arus. Data ini penting karena lokasi penelitian berada di wilayah pesisir, sehingga kondisi perairan sangat memengaruhi kelayakan aktivitas memancing."),
    ]:
        heading(title, 4)
        add_paragraph(text)
    add_table("Penilaian Weather Score", ["No", "Kondisi Cuaca", "Weather Score Dasar"], [[1, "Clear", 100], [2, "Clouds", 88], [3, "Mist, Fog, Haze", 70], [4, "Drizzle", 55], [5, "Rain", 35], [6, "Thunderstorm", 10]])
    add_code("Marine Score = (0.40 × Current Wave) + (0.25 × Max Wave 24 Jam) + (0.20 × Wave Stability) + (0.15 × Current Velocity)")
    add_image("dashboard-wave-opt-final.png", "Tampilan informasi gelombang dan kondisi perairan pada dashboard")
    heading("4.3.6.4 Fish Activity Engine", 4)
    add_paragraph("Fish Activity Engine digunakan untuk memperkirakan tingkat aktivitas ikan. Engine ini menggabungkan Solunar Score, Weather Score khusus aktivitas ikan, Pressure Score, dan Water Movement Score. Pada implementasi terbaru, BMKG tidak lagi digunakan sebagai parameter pembatas Fish Activity. BMKG tetap digunakan sebagai sumber prakiraan dan peringatan resmi, tetapi aktivitas ikan dihitung berdasarkan data live seperti kondisi cuaca, tekanan udara, angin live, gelombang, arus, dan suhu laut.")
    add_code("Fish Activity = (0.35 × S) + (0.15 × W) + (0.25 × P) + (0.25 × M) + Bonus - Penalti")
    add_paragraph("Bobot Solunar sebesar 35 persen digunakan karena waktu biologis ikan tetap menjadi faktor penting. Weather Score sebesar 15 persen digunakan untuk menggambarkan pengaruh kondisi langit terhadap aktivitas ikan. Pressure Score sebesar 25 persen digunakan karena tekanan udara berkaitan dengan perubahan kondisi lingkungan yang dapat memengaruhi perilaku ikan. Water Movement Score sebesar 25 persen digunakan karena arus dan gelombang rendah membantu pergerakan oksigen serta pakan alami.")
    add_table("Hasil Simulasi Fish Activity Engine", ["No", "Skenario", "Solunar Score", "Data Lingkungan", "Hasil"], [[1, "Di luar jam makan utama", 20, "Scattered clouds, 1008 hPa, angin 24,2 km/jam, gelombang 0,6 m, arus 0,9 m/s", "51,45%"], [2, "Periode minor", 60, "Scattered clouds, 1008 hPa, angin 24,2 km/jam, gelombang 0,6 m, arus 0,9 m/s", "65,45%"], [3, "Periode mayor", 90, "Scattered clouds, 1008 hPa, angin 24,2 km/jam, gelombang 0,6 m, arus 0,9 m/s", "75,95%"]])
    for title, text in [
        ("4.3.6.5 Spot Quality Engine", "Spot Quality Engine menghitung kualitas lokasi berdasarkan rating spot. Rating diperoleh dari data ulasan atau penilaian pengguna yang tersimpan pada Firestore. Semakin tinggi rating suatu spot, semakin tinggi pula Spot Quality Score yang diberikan."),
        ("4.3.6.6 User Preference Engine", "User Preference Engine digunakan untuk menyesuaikan hasil rekomendasi dengan kecenderungan pengguna. Nilai preferensi dapat dikaitkan dengan riwayat favorite, interaksi pengguna terhadap spot, atau kecenderungan pengguna memilih jenis spot tertentu."),
        ("4.3.6.7 Safety Engine", "Safety Engine merupakan lapisan koreksi yang digunakan untuk menilai keamanan aktivitas memancing. Berbeda dengan Fish Activity Engine, Safety Engine tetap mempertimbangkan BMKG sebagai sumber peringatan resmi. Hal ini dilakukan karena aspek keselamatan perlu dinilai lebih konservatif."),
    ]:
        heading(title, 4)
        add_paragraph(text)
    add_table("Penilaian Keamanan Angin pada Safety Engine", ["No", "Kecepatan Angin", "Kategori", "Nilai Safety"], [[1, "≤ 8 km/jam", "Tenang sampai sepoi ringan", "1,00"], [2, "9-14 km/jam", "Sepoi lembut", "0,95"], [3, "15-19 km/jam", "Cukup aman", "0,85"], [4, "20-28 km/jam", "Waspada", "0,70"], [5, "29-38 km/jam", "Berisiko untuk perahu kecil", "0,55"], [6, "39-49 km/jam", "Sangat berisiko", "0,40"], [7, "> 49 km/jam", "Tidak disarankan", "0,25"]])
    heading("4.3.6.8 Final Recommendation", 4)
    add_paragraph("Final Recommendation merupakan hasil akhir dari seluruh proses perhitungan. Skor dasar yang berasal dari Distance Engine, Weather Engine, Marine Engine, Fish Activity Engine, Spot Quality Engine, dan User Preference Engine dikalikan dengan Safety Factor. Setelah itu, sistem memberikan label rekomendasi berdasarkan rentang skor akhir.")
    add_table("Kategori Final Recommendation", ["No", "Rentang Skor", "Label Rekomendasi"], [[1, "≥ 85", "Sangat Direkomendasikan"], [2, "≥ 70 sampai < 85", "Direkomendasikan"], [3, "≥ 55 sampai < 70", "Cukup Layak"], [4, "≥ 40 sampai < 55", "Perlu Waspada"], [5, "< 40", "Tidak Direkomendasikan"]])


def add_remaining():
    heading("4.3.7 Community", 3)
    add_paragraph("Modul Community diimplementasikan sebagai media berbagi pengalaman antar pengguna. Fitur ini memungkinkan pengguna membuat postingan berisi foto, deskripsi, jenis ikan secara opsional, dan lokasi secara opsional. Implementasi dilakukan melalui CommunityFragment, CreatePostFragment, CommunityViewModel, PostAdapter, dan CommunityRepository.")
    add_paragraph("Pada proses pembuatan postingan, pengguna memilih foto dari galeri, kemudian foto tersebut diunggah ke Cloudinary. Setelah Cloudinary mengembalikan URL gambar, CommunityRepository menyimpan data postingan ke Cloud Firestore. Data yang tersimpan meliputi identitas pengguna, caption, URL gambar, jenis ikan, lokasi, jumlah like, jumlah komentar, dan jumlah favorite.")
    add_image("uiux-community.png", "Tampilan halaman Community aplikasi Fishing Point")
    heading("4.3.8 Profile", 3)
    add_paragraph("Modul Profile digunakan untuk menampilkan dan mengelola data pengguna. Implementasi dilakukan melalui ProfileFragment, ProfileViewModel, dan ProfileRepository. Halaman profile menampilkan nama pengguna, email, foto profil, statistik postingan, jumlah spot yang dibuat, jumlah favorite, daftar spot saya, daftar spot favorit, daftar postingan saya, reset password, dan logout.")
    add_image("uiux-profile.png", "Tampilan halaman Profile aplikasi Fishing Point")
    heading("4.3.9 Firestore", 3)
    add_paragraph("Cloud Firestore digunakan sebagai basis data utama aplikasi Fishing Point. Data yang disimpan meliputi data pengguna, fishing point, postingan community, komentar, favorite, review, notifikasi, serta cache data cuaca dan maritim. Firestore dipilih karena mendukung penyimpanan dokumen berbasis cloud dan sinkronisasi data secara realtime.")
    add_code("users\nfishing_points\ncommunity_posts\ncomments\nfavorites\nreviews\nnotifications\nweather_cache / bmkg_cache")
    heading("4.3.10 Cloudinary", 3)
    add_paragraph("Cloudinary digunakan sebagai layanan penyimpanan media. Hal ini dilakukan karena aplikasi menggunakan Firestore sebagai basis data, sedangkan file gambar membutuhkan layanan penyimpanan khusus. Cloudinary digunakan untuk menyimpan foto profil, foto spot, dan foto postingan komunitas. Setelah upload berhasil, Cloudinary mengembalikan secure URL yang kemudian disimpan pada Firestore.")
    heading("4.4 Pengujian")
    add_paragraph("Tahap pengujian dilakukan untuk memastikan bahwa aplikasi berjalan sesuai kebutuhan. Pengujian dilakukan melalui black box testing, unit testing, dan pengujian pada perangkat Android nyata. Black box testing digunakan untuk memeriksa fungsi dari sisi pengguna, unit testing digunakan untuk memvalidasi logika perhitungan, sedangkan pengujian perangkat nyata digunakan untuk memastikan aplikasi dapat berjalan pada lingkungan sebenarnya.")
    add_table("Ringkasan Pengujian Aplikasi", ["No", "Modul", "Skenario Pengujian", "Hasil"], [[1, "Authentication", "Login, register, session, reset password", "Berhasil"], [2, "Dashboard", "Menampilkan cuaca, perairan, rekomendasi, safety, dan fish activity", "Berhasil"], [3, "Maps", "Menampilkan marker, posisi pengguna, polyline, dan detail spot", "Berhasil"], [4, "Haversine", "Menghitung jarak antar koordinat", "Berhasil"], [5, "Recommendation Engine", "Menghasilkan skor rekomendasi dan label", "Berhasil"], [6, "Fish Activity Engine", "Skenario solunar rendah, sedang, dan tinggi", "Berhasil"], [7, "Community", "Upload foto, caption, like, komentar, favorite, share", "Berhasil"], [8, "Profile", "Edit profile, foto profile, statistik, reset password", "Berhasil"], [9, "Cloudinary", "Upload media dan menyimpan URL ke Firestore", "Berhasil"], [10, "Device nyata", "Aplikasi berjalan tanpa force close pada alur utama", "Berhasil"]])
    add_image("runtime-audit-map.png", "Pengujian runtime halaman Maps pada perangkat nyata")
    heading("4.5 Pembahasan")
    add_paragraph("Berdasarkan hasil implementasi dan pengujian, aplikasi Fishing Point telah memenuhi tujuan utama penelitian. Aplikasi mampu menampilkan lokasi fishing point pada peta, menghitung jarak menggunakan metode Haversine, menampilkan informasi cuaca dan perairan, serta menghasilkan rekomendasi memancing melalui Recommendation Engine. Fitur tambahan seperti Community dan Profile juga berhasil diimplementasikan sehingga aplikasi tidak hanya berfungsi sebagai pencari lokasi, tetapi juga sebagai media berbagi informasi antar pengguna.")
    add_paragraph("Metode Haversine berhasil diterapkan sebagai dasar perhitungan jarak. Hasil jarak digunakan pada beberapa bagian aplikasi, yaitu dashboard, maps, detail spot, daftar spot, dan Distance Engine. Dengan demikian, metode Haversine tidak berdiri sendiri, tetapi menjadi bagian dari alur rekomendasi yang lebih luas.")
    add_paragraph("Recommendation Engine berhasil menggabungkan beberapa parameter yang relevan dengan aktivitas memancing. Penggunaan OpenWeather dan Open-Meteo Marine membuat aplikasi dapat membaca kondisi live, sedangkan BMKG tetap digunakan sebagai prakiraan dan peringatan resmi. Pemisahan fungsi BMKG pada Fish Activity dan Safety membuat sistem menjadi lebih seimbang. Fish Activity menggunakan data live agar lebih sesuai dengan kondisi saat ini, sedangkan Safety Engine tetap mempertimbangkan BMKG agar penilaian keselamatan lebih berhati-hati.")
    add_paragraph("Secara keseluruhan, aplikasi Fishing Point dapat digunakan sebagai alat bantu pengambilan keputusan bagi pemancing. Aplikasi tidak menggantikan pertimbangan langsung di lapangan, tetapi memberikan informasi awal yang lebih terstruktur melalui kombinasi lokasi, cuaca, perairan, rekomendasi, dan keselamatan. Dengan hasil tersebut, implementasi aplikasi telah selaras dengan rumusan masalah dan tujuan penelitian yang dijelaskan pada bab sebelumnya.")


add_intro()
add_auth_dashboard_maps()
add_detail_haversine()
add_recommendation()
add_remaining()
doc.save(OUT)
print(OUT)
