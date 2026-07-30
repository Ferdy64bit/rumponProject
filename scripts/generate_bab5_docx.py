from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_doc_defaults(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)


def add_para(doc, text, bold=False, italic=False, align=None, first_line_cm=1.25, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(12)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(12)
    return p


doc = Document()
set_doc_defaults(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BAB V\nPENUTUP')
run.bold = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run.font.size = Pt(14)

add_heading(doc, '5.1 Kesimpulan')
add_para(doc, 'Berdasarkan hasil perancangan, implementasi, pengujian, dan pembahasan yang telah dilakukan pada penelitian Rancang Bangun Aplikasi Fishing Point Berbasis Android Menggunakan Metode Perhitungan Jarak Haversine di Desa Tanjung Anom, dapat disimpulkan bahwa aplikasi yang dibangun berhasil memenuhi tujuan utama penelitian. Aplikasi Fishing Point mampu membantu pengguna memperoleh informasi spot memancing secara lebih terarah melalui integrasi lokasi pengguna, perhitungan jarak Haversine, data cuaca live, data perairan, dan sistem rekomendasi yang dikembangkan secara bertahap sesuai kebutuhan lapangan.')
add_para(doc, 'Penerapan metode Haversine pada aplikasi terbukti menjadi dasar yang konsisten dalam menampilkan jarak antara posisi pengguna dan fishing point. Hasil perhitungan jarak tersebut tidak hanya digunakan untuk menampilkan daftar spot terdekat, tetapi juga menjadi salah satu komponen penting dalam pembentukan skor rekomendasi. Dengan demikian, pengguna dapat melihat spot yang lebih dekat sekaligus menilai apakah spot tersebut layak dipertimbangkan berdasarkan kondisi lingkungan yang sedang terjadi.')
add_para(doc, 'Selain perhitungan jarak, aplikasi juga berhasil mengintegrasikan berbagai sumber data eksternal sebagai pendukung pengambilan keputusan. Data cuaca dari OpenWeather, data perairan dari Open-Meteo Marine, dan informasi resmi dari BMKG digunakan untuk menyusun gambaran kondisi lapangan yang lebih utuh. Integrasi tersebut membuat aplikasi tidak hanya menampilkan lokasi, tetapi juga membantu pengguna memahami apakah situasi cuaca, gelombang, arus, dan kondisi laut mendukung aktivitas memancing pada saat tertentu.')
add_para(doc, 'Recommendation Engine yang dibangun pada aplikasi juga menunjukkan bahwa sistem dapat mengolah beberapa parameter secara bersamaan, seperti distance score, weather score, marine score, fish activity score, spot quality score, user preference score, dan safety multiplier. Penggabungan parameter tersebut menghasilkan rekomendasi yang lebih kontekstual dibandingkan hanya menampilkan spot berdasarkan kedekatan lokasi. Pada proses pengembangan, beberapa penyesuaian juga dilakukan agar hasil skor lebih sesuai dengan kondisi lapangan, terutama pada parameter keamanan dan aktivitas ikan yang memerlukan kalibrasi berdasarkan observasi nyata.')
add_para(doc, 'Dari sisi antarmuka dan pengalaman pengguna, aplikasi Fishing Point telah menyediakan beberapa modul utama yang saling terhubung, yaitu authentication, dashboard, maps, detail spot, community, profile, Firestore, dan Cloudinary. Modul dashboard berfungsi sebagai pusat informasi utama, sedangkan modul maps digunakan untuk menampilkan marker spot, polyline navigasi, dan informasi spot terdekat. Modul detail spot menampilkan data yang lebih lengkap berdasarkan koordinat spot, sementara modul community dan profile mendukung interaksi sosial serta pengelolaan data pengguna. Seluruh modul tersebut dirancang dengan pendekatan MVVM dan repository pattern agar struktur aplikasi tetap terjaga, mudah dipelihara, dan mudah dikembangkan di kemudian hari.')
add_para(doc, 'Pengujian yang dilakukan pada perangkat nyata memperlihatkan bahwa alur utama aplikasi telah berjalan dengan baik. Fitur login, register, reset password, tampilan dashboard, maps, detail spot, community, profile, serta proses upload media telah dapat digunakan sesuai fungsinya. Beberapa penyempurnaan pada tampilan juga berhasil meningkatkan stabilitas aplikasi pada berbagai ukuran perangkat. Hal ini menunjukkan bahwa aplikasi tidak hanya selesai pada tahap rancangan, tetapi benar-benar dapat dioperasikan sebagai produk Android yang relevan untuk mendukung kebutuhan pengguna dalam mencari dan menilai titik memancing.')
add_para(doc, 'Secara keseluruhan, penelitian ini membuktikan bahwa kombinasi metode Haversine, integrasi data cuaca dan perairan, serta pengembangan recommendation engine dapat diterapkan untuk membangun aplikasi Fishing Point yang fungsional dan bermanfaat. Aplikasi ini masih memiliki ruang pengembangan, namun hasil implementasi yang diperoleh sudah cukup kuat untuk mendukung kebutuhan penelitian dan menjadi dasar bagi pengembangan lanjutan pada tahap berikutnya.')

add_heading(doc, '5.2 Saran')
add_para(doc, 'Berdasarkan hasil implementasi dan pengujian yang telah dilakukan, terdapat beberapa saran yang dapat digunakan untuk pengembangan aplikasi Fishing Point pada tahap selanjutnya. Pertama, aplikasi dapat dikembangkan dengan menambahkan notifikasi kondisi perairan ekstrem atau peringatan cuaca tertentu agar pengguna memperoleh informasi yang lebih proaktif sebelum berangkat ke lokasi memancing. Fitur ini akan memberikan nilai tambah karena aplikasi tidak hanya bersifat informatif, tetapi juga preventif terhadap kondisi yang berisiko.')
add_para(doc, 'Kedua, pengembangan berikutnya dapat difokuskan pada penyempurnaan model rekomendasi agar lebih adaptif terhadap kondisi lapangan yang berubah cepat. Meskipun engine yang saat ini digunakan sudah memadukan beberapa parameter utama, akurasi sistem masih dapat ditingkatkan melalui penyesuaian bobot, kalibrasi berdasarkan data lapangan yang lebih panjang, serta evaluasi berkala terhadap hasil rekomendasi yang muncul pada pengguna.')
add_para(doc, 'Ketiga, modul komunitas dan profile masih dapat diperluas agar aktivitas pengguna semakin lengkap. Pengembangan lanjutan dapat mencakup penyaringan postingan yang lebih baik, rekap statistik pengguna, manajemen media yang lebih fleksibel, serta peningkatan fitur personalisasi pada spot favorit dan spot pribadi. Dengan pengembangan tersebut, aplikasi dapat menjadi lebih interaktif sekaligus lebih sesuai dengan kebutuhan pengguna yang aktif berbagi pengalaman memancing.')
add_para(doc, 'Keempat, dari sisi teknis, aplikasi masih dapat dikembangkan agar performanya lebih ringan ketika memuat data dalam jumlah besar. Optimasi cache, pengelolaan query Firestore, pembatasan request API, dan pengaturan refresh data dapat dipertimbangkan agar aplikasi tetap stabil ketika digunakan secara berulang pada perangkat yang berbeda. Hal ini penting karena aplikasi berbasis lokasi dan data live cenderung memerlukan sinkronisasi data yang cukup sering.')
add_para(doc, 'Kelima, penelitian selanjutnya dapat memperluas wilayah implementasi aplikasi ke daerah pesisir lain agar model perhitungan yang digunakan semakin teruji pada berbagai karakteristik perairan. Dengan cakupan data yang lebih luas, aplikasi Fishing Point berpotensi menjadi sistem pendukung keputusan yang lebih matang bagi pengguna di wilayah pesisir lainnya, bukan hanya terbatas pada area penelitian awal.')
add_para(doc, 'Dengan demikian, saran-saran tersebut diharapkan dapat menjadi acuan apabila aplikasi Fishing Point dikembangkan kembali pada versi berikutnya. Penyempurnaan yang dilakukan secara bertahap akan membuat sistem semakin stabil, relevan, dan sesuai dengan kebutuhan pengguna di lapangan.')

add_heading(doc, '5.3 Penutup')
add_para(doc, 'BAB V ini disusun untuk merangkum hasil akhir dari proses penelitian dan pengembangan aplikasi Fishing Point. Hasil yang diperoleh menunjukkan bahwa penelitian telah berhasil menghasilkan aplikasi Android yang menggabungkan pencarian titik memancing, perhitungan jarak, integrasi data lingkungan, recommendation engine, community, dan profile dalam satu sistem yang terintegrasi. Dengan hasil tersebut, aplikasi Fishing Point dapat dijadikan sebagai solusi awal yang membantu pengguna menentukan spot memancing secara lebih terarah dan informatif.')

out = r'C:\Users\ThinkPad T14\Music\BAB V - Penutup Fishing Point.docx'
doc.save(out)
print(out)
